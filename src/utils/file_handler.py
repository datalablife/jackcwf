"""File handling utilities for document processing."""

import io
import logging
from typing import Optional
from pathlib import Path

from fastapi import UploadFile, HTTPException, status

logger = logging.getLogger(__name__)


class FileHandler:
    """
    Handler for processing uploaded files.

    Supports:
    - PDF files (.pdf)
    - Text files (.txt, .md)
    - Word documents (.docx)
    - CSV files (.csv)
    """

    # File size limits
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    # Supported file types
    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".txt": "txt",
        ".md": "markdown",
        ".docx": "docx",
        ".csv": "csv",
    }

    SUPPORTED_MIME_TYPES = {
        "application/pdf": "pdf",
        "text/plain": "txt",
        "text/markdown": "markdown",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/csv": "csv",
    }

    def detect_file_type(self, filename: str) -> str:
        """
        Detect file type from filename.

        Args:
            filename: Name of the file

        Returns:
            File type (pdf, txt, markdown, docx, csv)

        Raises:
            ValueError: If file type is not supported
        """
        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}"
            )

        return self.SUPPORTED_EXTENSIONS[ext]

    async def extract_text(self, file: UploadFile) -> str:
        """
        Extract text content from uploaded file.

        Args:
            file: Uploaded file

        Returns:
            Extracted text content

        Raises:
            HTTPException: If extraction fails
        """
        file_type = self.detect_file_type(file.filename)

        try:
            # Read file content
            content = await file.read()

            # Reset file pointer
            await file.seek(0)

            # Extract based on file type
            if file_type == "txt" or file_type == "markdown":
                return self._extract_text_plain(content)
            elif file_type == "pdf":
                return await self._extract_text_pdf(content)
            elif file_type == "docx":
                return await self._extract_text_docx(content)
            elif file_type == "csv":
                return self._extract_text_csv(content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error extracting text from {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract text from file: {str(e)}",
            )

    def _extract_text_plain(self, content: bytes) -> str:
        """
        Extract text from plain text file.

        Args:
            content: File content as bytes

        Returns:
            Decoded text content
        """
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue

        # If all fail, use utf-8 with error handling
        return content.decode('utf-8', errors='replace')

    async def _extract_text_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF file.

        Args:
            content: PDF file content as bytes

        Returns:
            Extracted text content
        """
        try:
            import pypdf

            # Create PDF reader
            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file)

            # Extract text from all pages
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)

            if not text_parts:
                raise ValueError("No text content found in PDF")

            return "\n\n".join(text_parts)

        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="PDF processing library not installed. Please install pypdf.",
            )
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    async def _extract_text_docx(self, content: bytes) -> str:
        """
        Extract text from Word document.

        Args:
            content: DOCX file content as bytes

        Returns:
            Extracted text content
        """
        try:
            import docx

            # Create document reader
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            if not text_parts:
                raise ValueError("No text content found in DOCX")

            return "\n\n".join(text_parts)

        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="DOCX processing library not installed. Please install python-docx.",
            )
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    def _extract_text_csv(self, content: bytes) -> str:
        """
        Extract text from CSV file.

        Args:
            content: CSV file content as bytes

        Returns:
            Formatted text content
        """
        import csv

        try:
            # Decode content
            text = self._extract_text_plain(content)

            # Parse CSV
            reader = csv.reader(io.StringIO(text))

            # Convert to formatted text
            lines = []
            for row in reader:
                if row:  # Skip empty rows
                    lines.append(" | ".join(row))

            if not lines:
                raise ValueError("No content found in CSV")

            return "\n".join(lines)

        except Exception as e:
            logger.error(f"Error extracting CSV text: {str(e)}")
            raise ValueError(f"Failed to extract text from CSV: {str(e)}")


def validate_file_upload(file: UploadFile) -> None:
    """
    Validate uploaded file.

    Args:
        file: Uploaded file

    Raises:
        HTTPException: If validation fails
    """
    # Check if file is provided
    if not file:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No file provided",
        )

    # Check filename
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid filename",
        )

    # Check file extension
    handler = FileHandler()
    try:
        file_type = handler.detect_file_type(file.filename)
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        )

    # Check MIME type if available
    if file.content_type:
        if file.content_type not in FileHandler.SUPPORTED_MIME_TYPES:
            logger.warning(
                f"MIME type {file.content_type} not in supported list, "
                f"but extension is valid. Proceeding with caution."
            )

    logger.info(f"File validation passed for {file.filename} (type: {file_type})")
